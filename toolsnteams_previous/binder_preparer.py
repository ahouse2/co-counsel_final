from docx import Document
from docx.shared import Inches
from neuro_san.interfaces.coded_tool import CodedTool

class BinderPreparer(CodedTool):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    def prepare_binder(self, filepath: str, evidence_list: list[dict], case_name: str = "Untitled Case"):
        """
        Prepares a trial binder document from a list of evidence items.

        :param filepath: The path to save the generated Word document.
        :param evidence_list: A list of dictionaries, each representing an evidence item.
                              Each dict should have at least 'name', 'type', and 'url'.
                              Optionally 'annotation' for notes.
        :param case_name: The name of the case for the binder title.
        """
        document = Document()

        # Title Page
        document.add_heading(f"Trial Binder for {case_name}", level=0)
        document.add_paragraph(f"Date Prepared: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_page_break()

        # Table of Contents (will be updated later)
        document.add_heading("Table of Contents", level=1)
        toc_paragraph = document.add_paragraph()
        document.add_page_break()

        # Evidence Sections
        for i, evidence_item in enumerate(evidence_list):
            document.add_heading(f"Exhibit {i+1}: {evidence_item.get('name', 'N/A')}", level=1)
            document.add_paragraph(f"Type: {evidence_item.get('type', 'N/A')}")
            document.add_paragraph(f"URL: {evidence_item.get('url', 'N/A')}")
            
            annotation = evidence_item.get('annotation')
            if annotation:
                document.add_paragraph(f"Annotation: {annotation}")
            
            document.add_paragraph("\n--- Evidence Content Placeholder ---")
            document.add_paragraph(f"Refer to original document at: {evidence_item.get('url', 'N/A')}")
            document.add_page_break()

        # Save the document
        document.save(filepath)
        print(f"Trial binder '{case_name}' prepared and saved to {filepath}")

from datetime import datetime
