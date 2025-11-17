from docx import Document
from docx.shared import Inches
from neuro_san.interfaces.coded_tool import CodedTool
from toolsnteams_previous.template_library import TemplateLibrary


class DocumentDrafter(CodedTool):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.template_library = TemplateLibrary()

    def create_document(self, filepath: str, content: str):
        """
        Creates a new Word document.

        :param filepath: The path to the new document.
        :param content: The content to add to the document.
        """
        document = Document()
        document.add_paragraph(content)
        document.save(filepath)

    def add_paragraph(self, filepath: str, content: str):
        """
        Adds a new paragraph to an existing Word document.

        :param filepath: The path to the document.
        :param content: The content to add.
        """
        document = Document(filepath)
        document.add_paragraph(content)
        document.save(filepath)

    def add_heading(self, filepath: str, text: str, level: int):
        """
        Adds a new heading to an existing Word document.

        :param filepath: The path to the document.
        :param text: The text of the heading.
        :param level: The level of the heading (1-9).
        """
        document = Document(filepath)
        document.add_heading(text, level=level)
        document.save(filepath)

    def add_numbered_list(self, filepath: str, items: list[str]):
        """
        Adds a numbered list to an existing Word document.

        :param filepath: The path to the document.
        :param items: A list of strings, where each string is an item in the numbered list.
        """
        document = Document(filepath)
        for item in items:
            document.add_paragraph(item, style='List Number')
        document.save(filepath)

    def add_bullet_list(self, filepath: str, items: list[str]):
        """
        Adds a bulleted list to an existing Word document.

        :param filepath: The path to the document.
        :param items: A list of strings, where each string is an item in the bulleted list.
        """
        document = Document(filepath)
        for item in items:
            document.add_paragraph(item, style='List Bullet')
        document.save(filepath)

    def add_table(self, filepath: str, data: list[list[str]], heading_row: bool = True):
        """
        Adds a table to an existing Word document.

        :param filepath: The path to the document.
        :param data: A list of lists, where each inner list represents a row of cells.
        :param heading_row: If True, the first row will be treated as a heading row.
        """
        document = Document(filepath)
        if not data:
            return

        rows = len(data)
        cols = len(data[0])
        table = document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        for r_idx, row_data in enumerate(data):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_data in enumerate(row_data):
                row_cells[c_idx].text = cell_data
            if r_idx == 0 and heading_row:
                row_cells[0].paragraphs[0].runs[0].bold = True # Make heading row bold

        document.save(filepath)

    def add_section(self, filepath: str, title: str, content: str):
        """
        Adds a new section with a title and content to an existing Word document.

        :param filepath: The path to the document.
        :param title: The title of the section.
        :param content: The content of the section.
        """
        document = Document(filepath)
        document.add_heading(title, level=2)
        document.add_paragraph(content)
        document.save(filepath)

    def draft_legal_document(self, filepath: str, motion_type: str, data: dict):
        """
        Drafts a legal document based on the specified motion type and provided data.

        :param filepath: The path to save the drafted document.
        :param motion_type: The type of motion to draft (e.g., "motion_to_dismiss").
        :param data: A dictionary containing data to populate the template,
                     e.g., {"facts": "...", "theories": "...", "conflicts": "..."}.
        """
        try:
            # Build the prompt using the TemplateLibrary
            # The TemplateLibrary's build_prompt method already formats the string
            # with facts, theories, and conflicts from its internal data.
            # We need to ensure that the 'data' dictionary passed here
            # can override or supplement the internal data if necessary,
            # or that the TemplateLibrary is initialized with the correct context.
            # For now, we'll assume TemplateLibrary handles data retrieval internally
            # and its build_prompt returns a fully formatted string.
            
            # If the data dict contains specific overrides for facts, theories, conflicts,
            # we might need to pass them to build_prompt or format the template here.
            # For simplicity, let's assume data dict directly provides the formatted strings
            # that would otherwise come from the DB via TemplateLibrary.
            
            # A more robust solution would involve passing the data to the TemplateLibrary
            # so it can intelligently merge/override its internal data.
            # For this iteration, we'll use the TemplateLibrary to get the base template
            # and then format it with the provided 'data' dictionary.
            
            template_string = self.template_library.MOTION_TEMPLATES.get(motion_type)
            if not template_string:
                raise ValueError(f"Unknown motion type: {motion_type}")

            # Format the template string with the provided data
            # This assumes 'data' keys match the template placeholders (facts, theories, conflicts)
            formatted_content = template_string.format(**data)

            # Create a new document and add the formatted content
            document = Document()
            document.add_paragraph(formatted_content)
            document.save(filepath)
            print(f"Successfully drafted '{motion_type}' to {filepath}")

        except ValueError as e:
            print(f"Error drafting document: {e}")
        except Exception as e:
            print(f"An unexpected error occurred: {e}")

